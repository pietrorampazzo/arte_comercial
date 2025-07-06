from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def criar_template_proposta(path_saida, dados, itens):
    doc = Document()

    # Estilo do documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Cabeçalho
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run(f"Dispensa Eletrônica Nº {dados['numero_pregao']} (Lei 14.133/2021)").bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(f"UASG {dados['uasg']} - {dados['orgao']}")

    doc.add_paragraph()

    # Tabela resumo
    resumo_headers = [
        "UASG", "PREGÃO", "Método de Envio", "Termos de Envio",
        "Data de Pregão", "Termos de Pagamento", "Data de Vencimento da proposta"
    ]
    resumo_valores = [
        dados['uasg'], dados['numero_pregao'], dados.get('metodo_envio', ''),
        dados.get('termos_envio', ''), dados['dia_pregao'],
        dados['termos_pagamento'], dados['vencimento_proposta']
    ]

    tabela_resumo = doc.add_table(rows=2, cols=len(resumo_headers))
    tabela_resumo.style = 'Table Grid'

    for i, header in enumerate(resumo_headers):
        tabela_resumo.cell(0, i).text = header
        tabela_resumo.cell(1, i).text = resumo_valores[i]

    doc.add_paragraph()

    # Tabela de itens
    tabela_itens = doc.add_table(rows=1, cols=6)
    tabela_itens.style = 'Table Grid'
    headers_itens = ["ITEM", "Marca", "Descrição", "Qnt.", "Valor", "Valor Total"]
    for i, header in enumerate(headers_itens):
        tabela_itens.cell(0, i).text = header

    for item in itens:
        row = tabela_itens.add_row().cells
        row[0].text = str(item['item'])
        row[1].text = item['marca']
        row[2].text = item['descricao']
        row[3].text = str(item['quantidade'])
        row[4].text = f"R$ {item['valor_unitario']:.2f}".replace('.', ',')
        row[5].text = f"R$ {item['valor_total']:.2f}".replace('.', ',')

    doc.add_paragraph()
    doc.add_paragraph(
        "Nos preços acima estão incluídas todas as despesas com impostos, fretes, taxas, "
        "descargas e quaisquer outras que incidam direta ou indiretamente no fornecimento dos materiais desta licitação. "
        "* A entrega dos materiais será feita no local determinado pelo Órgão Licitante "
        "** O prazo de validade da proposta não será inferior a 60 (sessenta) dias, a contar da data de sua apresentação."
    )

    doc.add_paragraph()
    doc.add_paragraph("Dados Bancários: Banco do Brasil (001) Agência: 4770-1  Conta Corrente: 106261-1")

    doc.save(path_saida)
    print(f"✅ Documento salvo em: {path_saida}")

# Exemplo de uso
dados_licitacao = {
    'uasg': '160411',
    'numero_pregao': '90031/2024',
    'orgao': '7 BATALHAO DE INFANTARIA BLINDADO/RS',
    'dia_pregao': '20/06',
    'termos_pagamento': '30 d.d.',
    'vencimento_proposta': '60 d.d.',
    'metodo_envio': 'TRANSPORTADORA',
    'termos_envio': ''
}

itens = [
    {
        'item': 13,
        'marca': 'Remo - Encore',
        'descricao': 'PELE ENCORE 14 POL AMBASSADOR POROSA EN-0114-BA RE',
        'quantidade': 2,
        'valor_unitario': 171.25,
        'valor_total': 342.50
    }
]

output = r"C:\Users\pietr\OneDrive\Área de Trabalho\ARTE\MODELO_PROPOSTA.docx"
criar_template_proposta(output, dados_licitacao, itens)
